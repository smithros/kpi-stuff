import string
from itertools import *

from docx import Document

alphabet_ukr = "АБВГҐДЕЄЖЗИІЇЙКЛМНОПРСТУФХЦЧШЩЬЮЯабвгґдеєжзиіїйклмнопрстуфхцчшщьюя1234567890 !?.,-"
alphabet_eng = string.ascii_letters + string.digits


def encrypt_start(key: str, key2: str, key3: str, text: str, lang: str) -> str:
    mod = get_vars(lang)[0]
    hex_int = get_vars(lang)[2]
    ord_msg = [ord(i) for i in text]
    enc_1 = encrypt(key, ord_msg, mod)
    enc_2 = encrypt(key2, enc_1, mod)
    enc_3 = encrypt(key3, enc_2, mod)
    result = []
    for i in enc_3:
        result.append(format(i, hex_int))
    return ''.join(result)


def encrypt(key: str, message, mod: int) -> list:
    ord_key = [ord(j) for j in key]
    arr = set_key(ord_key, mod)
    stream = get_stream(arr, mod)
    result = []
    for k in message:
        tmp = next(stream)
        result.append(k ^ tmp)
    return result


def decrypt_start(key: str, key2: str, key3: str, message: str, lang: str) -> str:
    mod = get_vars(lang)[0]
    step = get_vars(lang)[1]
    ord_message = [int(message[i:i + step], 16) for i in range(0, len(message), step)]
    dec_1 = decrypt(key3, ord_message, mod)
    dec_2 = decrypt(key2, dec_1, mod)
    dec_3 = decrypt(key, dec_2, mod)
    result = []
    for i in dec_3:
        result.append(chr(i))
    return ''.join(result)


def decrypt(key: str, message, mod: int) -> list:
    ord_key = [ord(j) for j in key]
    arr = set_key(ord_key, mod)
    stream = get_stream(arr, mod)
    result = []
    for k in message:
        tmp = next(stream)
        result.append(k ^ tmp)
    return result


def get_vars(lang: str) -> list:
    mod = 0
    step = 0
    hex_int = ''
    if lang == 'ukr':
        mod = 2048
        hex_int = '03X'
        step = 3
    elif lang == 'eng':
        mod = 256
        hex_int = '02X'
        step = 2
    return [mod, step, hex_int]


def set_key(ord_key: list, mod: int) -> list:
    arr = list(range(mod))
    j = 0
    for i in range(mod):
        j = (j + arr[i] + ord_key[i % len(ord_key)]) % mod
        arr[i], arr[j] = arr[j], arr[i]
    return arr


def get_stream(arr: list, mod: int):
    i = 0
    j = 0
    while True:
        i = (i + 1) % mod
        j = (j + arr[i]) % mod
        arr[i], arr[j] = arr[j], arr[i]
        yield arr[(arr[i] + arr[j]) % mod]


def brute(message: str, lang: str):
    keys1, keys2 = [], []
    mod = get_vars(lang)[0]
    step = get_vars(lang)[1]
    ord_message = [int(message[i:i + step], 16) for i in range(0, len(message), step)]
    for q in range(1, len(ord_message) + 1):
        for k in range(1, len(ord_message) + 1):
            if lang == 'ukr':
                keys1 = product(alphabet_ukr, repeat=q)
                keys2 = product(alphabet_ukr, repeat=k)
            elif lang == 'eng':
                keys1 = product(alphabet_eng, repeat=q)
                keys2 = product(alphabet_eng, repeat=k)
            for key1 in keys1:
                for key2 in keys2:
                    ord_key1 = [ord(j) for j in key1]
                    arr1 = set_key(ord_key1, mod)
                    stream1 = get_stream(arr1, mod)
                    result = []
                    for j in ord_message:
                        tmp = next(stream1)
                        result.append((j ^ tmp))
                    ord_key2 = [ord(j) for j in key2]
                    arr2 = set_key(ord_key2, mod)
                    stream2 = get_stream(arr2, mod)
                    translated = []
                    for c in result:
                        tmp = next(stream2)
                        translated.append(chr(c ^ tmp))


def read_doc(path: str) -> str:
    doc = Document(path)
    tmp = [p.text for p in doc.paragraphs]
    res = "".join(tmp)
    return res


def write_doc(path: str, text: str) -> str:
    doc = Document()
    doc.add_paragraph(text)
    doc.save(path)
    return text


def read_txt(file_name: str) -> str:
    in_d_ukr = open(file_name, 'r')
    text_d_ukr = in_d_ukr.read()
    in_d_ukr.close()
    return text_d_ukr


def write_txt(file_name: str, text: str):
    out_dec_ukr = open(file_name, encoding='utf-8', mode='w')
    out_dec_ukr.write(text)
    out_dec_ukr.close()


def print_lang_test(msg: str, k1: str, k2: str, k3: str, lang: str):
    print("\nMessage: ", msg)
    print(f"Keys: {k1} {k2} {k3}")
    encrypt_text_ukr = encrypt_start(k1, k2, k3, msg, lang)
    print("Encrypted message: ", encrypt_text_ukr)
    decrypt_text_ukr = decrypt_start(k1, k2, k3, encrypt_text_ukr, lang)
    print("Decrypted message: ", decrypt_text_ukr)


def print_file_test(e_msg, e_text, d_msg, d_text, k1: str, k2: str, k3: str):
    print('\nMessage in ukr file to encrypt: ', e_msg)
    print(f"Keys: {k1} {k2} {k3}")
    print('Encryption result: ', e_text)
    print('\nMessage in ukr file to decrypt: ', d_msg)
    print(f"Keys: {k1} {k2} {k3}")
    print('Decryption result: ', d_text)


def get_la6_results():
    ukr_msg = 'Захист інформації у компютерних системах.'
    eng_msg = 'Information security work'
    key_ukr, key_ukr2, key_ukr3 = 'один', 'два', 'три'
    key_eng, key_eng2, key_eng3 = 'one', 'two', 'three'

    print_lang_test(ukr_msg, key_ukr, key_ukr2, key_ukr, "ukr")
    print_lang_test(eng_msg, key_eng, key_eng2, key_eng3, "ukr")

    in_enc_ukr = 'in_enc_ukr.txt'
    in_dec_ukr = 'in_dec_ukr.txt'
    out_enc_ukr = 'out_enc_ukr.txt'
    out_dec_ukr = 'out_dec_ukr.txt'
    in_enc_eng = 'in_enc_eng.docx'
    in_dec_eng = 'in_dec_eng.docx'
    out_enc_eng = 'out_enc_eng.docx'
    out_dec_eng = 'out_dec_eng.docx'

    text_e_ukr = read_txt(in_enc_ukr)
    text_d_ukr = read_txt(in_dec_ukr)

    enc_text_ukr = encrypt_start(key_ukr, key_ukr2, key_ukr3, text_e_ukr, 'ukr')
    write_txt(out_enc_ukr, enc_text_ukr)

    dec_text_ukr = decrypt_start(key_ukr, key_ukr2, key_ukr3, text_d_ukr, 'ukr')
    write_txt(out_dec_ukr, dec_text_ukr)

    enc_text_eng = read_doc(in_enc_eng)
    text_encr_eng = write_doc(out_enc_eng, encrypt_start(key_eng, key_eng2, key_eng3, enc_text_eng, 'eng'))
    dec_text_eng = read_doc(in_dec_eng)
    text_decr_eng = write_doc(out_dec_eng, decrypt_start(key_eng, key_eng2, key_eng3, dec_text_eng, 'eng'))

    print_file_test(text_e_ukr, enc_text_ukr, text_d_ukr, dec_text_ukr, key_ukr, key_ukr2, key_ukr3)
    print_file_test(enc_text_eng, text_encr_eng, dec_text_eng, text_decr_eng, key_eng, key_eng2, key_eng3)


if __name__ == '__main__':
    get_la6_results()
