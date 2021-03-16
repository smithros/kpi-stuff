import random

from math import gcd
from docx import Document

ukr_symbols = "ЙЦУКЕНГШЩЗХЇФІВАПРОЛДЖЄЯЧСМИТЬБЮйцукенгшщзхїфівапролджєячсмитьбю1234567890 :.,’"


def decrypt(text_to_dec: str, num: int) -> str:
    k1, k2 = num // len(ukr_symbols), num % len(ukr_symbols)
    k3 = pow(k1, -1, len(ukr_symbols))
    res = []
    for i in text_to_dec:
        res.append(ukr_symbols[((ukr_symbols.index(i) - k2) * k3) % len(ukr_symbols)])
    return ''.join(res)


def encrypt(text_to_enc: str, num: int) -> str:
    k1, k2 = num // len(ukr_symbols), num % len(ukr_symbols)
    res = []
    for i in text_to_enc:
        res.append(ukr_symbols[(k1 * ukr_symbols.index(i) + k2) % len(ukr_symbols)])
    return ''.join(res)


def generate_rand_key() -> int:
    while True:
        first_key = random.randint(2, len(ukr_symbols))
        second_key = random.randint(2, len(ukr_symbols))
        if gcd(second_key, len(ukr_symbols)) == 1:
            return second_key * len(ukr_symbols) + first_key


def get_all_keys() -> list:
    keys, res = [], []
    for i in range(2, len(ukr_symbols)):
        if gcd(len(ukr_symbols), i) == 1:
            res.append(i)
    for j in res:
        for k in range(2, len(ukr_symbols)):
            keys.append(j * len(ukr_symbols) + k)
    return keys


def brute(text: str) -> list:
    keys, hack_res = get_all_keys(), []
    for ky in keys:
        k1, k2 = ky // len(ukr_symbols), ky % len(ukr_symbols)
        k3 = pow(k1, -1, len(ukr_symbols))
        hacked_text = []
        for i in text:
            hacked_text.append(ukr_symbols[((ukr_symbols.index(i) - k2) * k3) % len(ukr_symbols)])
        hack_res.append("key = " + str(ky) + ' text: ' + ''.join(hacked_text))
    return hack_res


def read_docx(path: str) -> str:
    doc = Document(path)
    tmp = [p.text for p in doc.paragraphs]
    res = "".join(tmp)
    return res


def write_docx(path: str, prg: str):
    doc = Document()
    doc.add_paragraph(prg)
    doc.save(path)
